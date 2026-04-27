"""
File loaders cho RAG — extract text từ PDF/DOCX/TXT/MD.

Mỗi loader trả về string (full text). Tách splitting/embedding ra layer khác.
"""
from __future__ import annotations
import logging
import re
from pathlib import Path
from typing import Callable, Dict

log = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


# ---------------------------------------------------------------------------
# Per-format loaders
# ---------------------------------------------------------------------------

def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:
            log.warning("PDF page %d extract failed: %s", i, e)
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Bảng (tables) cũng cần lấy
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


LOADERS: Dict[str, Callable[[Path], str]] = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".txt":  _load_text,
    ".md":   _load_text,
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

class UnsupportedFileType(ValueError):
    """File extension không nằm trong SUPPORTED_EXTS."""


def detect_type(filename: str) -> str:
    """Trả về extension dạng .ext (lowercase). Raise nếu không hỗ trợ."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise UnsupportedFileType(
            f"Định dạng {ext or '(không có đuôi)'} không được hỗ trợ. "
            f"Chỉ chấp nhận: {', '.join(sorted(SUPPORTED_EXTS))}"
        )
    return ext


def extract_text(path: Path | str) -> str:
    """Đọc file → trả về plain text đã làm sạch."""
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"File không tồn tại: {p}")

    ext = detect_type(p.name)
    raw = LOADERS[ext](p)
    return clean_text(raw)


def clean_text(text: str) -> str:
    """Bỏ ký tự thừa, gộp khoảng trắng — không phá cấu trúc đoạn."""
    if not text:
        return ""
    # Chuẩn hóa newline
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Bỏ các khoảng trắng trùng (giữ \n)
    text = re.sub(r"[ \t]+", " ", text)
    # Bỏ blank line dư (giữ tối đa 2 newline liên tiếp)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
